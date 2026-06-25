#!/usr/bin/env python3
"""把缺少语义样式的 DOCX 导入为可重建的 Markdown 来源层。"""

from __future__ import annotations

import hashlib
import json
import re
import shutil
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as DocumentType
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
OUT_ROOT = ROOT / "sources" / "imported"
DOCS = {
    "agent-from-zero": ROOT / "大模型AI Agent知识从0-1笔记-万字详解版本！ .docx",
    "juliet-llm": ROOT / "居丽叶LLM体系知识搭建.docx",
}

DECIMAL_HEADING = re.compile(
    r"^(?P<num>\d+(?:\.\d+)+)[\s、.]+(?P<title>\S.*)$"
)
DOT_TOP_HEADING = re.compile(r"^(?P<num>\d+)\.\s+(?P<title>\S.*)$")
CHAPTER_TOP_HEADING = re.compile(r"^(?P<num>\d+)\s+\S.*篇$")
CHINESE_HEADING = re.compile(r"^第?[一二三四五六七八九十百]+[章节篇、.]\s*\S+")


@dataclass
class Block:
    kind: str
    text: str
    level: int | None = None
    section: str | None = None
    images: list[str] | None = None


def slugify(text: str) -> str:
    text = re.sub(r"[^\w\u4e00-\u9fff-]+", "-", text.lower())
    return re.sub(r"-+", "-", text).strip("-")[:80] or "section"


def clean(text: str) -> str:
    return re.sub(r"[ \t\xa0]+", " ", text.replace("\r", "")).strip()


def heading_info(text: str, source_id: str) -> tuple[int, str | None] | None:
    text = clean(text)
    match = DECIMAL_HEADING.match(text)
    if match:
        number = match.group("num")
        return min(number.count(".") + 1, 5), number
    if source_id == "agent-from-zero":
        match = DOT_TOP_HEADING.match(text)
        if match:
            return 1, match.group("num")
    if source_id == "juliet-llm":
        match = CHAPTER_TOP_HEADING.match(text)
        if match:
            return 1, match.group("num")
    if CHINESE_HEADING.match(text):
        return 2, None
    return None


def paragraph_images(paragraph: Paragraph, rels, media_dir: Path) -> list[str]:
    found: list[str] = []
    for blip in paragraph._p.xpath(".//*[local-name()='blip']"):
        rid = blip.get(
            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
        )
        if not rid or rid not in rels:
            continue
        part = rels[rid].target_part
        digest = hashlib.sha1(part.blob).hexdigest()[:12]
        suffix = Path(str(part.partname)).suffix.lower() or ".bin"
        name = f"{digest}{suffix}"
        target = media_dir / name
        if not target.exists():
            target.write_bytes(part.blob)
        found.append(name)
    return found


def paragraph_block(
    paragraph: Paragraph, rels, media_dir: Path, source_id: str
) -> Block | None:
    text = clean(paragraph.text)
    images = paragraph_images(paragraph, rels, media_dir)
    if not text and not images:
        return None
    info = heading_info(text, source_id) if text else None
    return Block(
        kind="paragraph",
        text=text,
        level=info[0] if info else None,
        section=info[1] if info else None,
        images=images,
    )


def cell_text(cell) -> str:
    return "<br>".join(clean(p.text) for p in cell.paragraphs if clean(p.text))


def table_block(table: Table) -> Block | None:
    rows = [[cell_text(cell) for cell in row.cells] for row in table.rows]
    rows = [row for row in rows if any(row)]
    if not rows:
        return None
    width = max(len(row) for row in rows)
    rows = [row + [""] * (width - len(row)) for row in rows]

    # 单列表格或超长单元格通常是从网页/飞书复制来的排版容器，应展开为正文。
    if width == 1 or max(len(cell) for row in rows for cell in row) > 1000:
        flattened = "\n\n".join(cell for row in rows for cell in row if cell)
        return Block(kind="layout-table", text=flattened)

    def esc(value: str) -> str:
        return value.replace("|", "\\|").replace("\n", "<br>")

    header = rows[0]
    body = rows[1:] or [[""] * width]
    lines = [
        "| " + " | ".join(esc(v) or " " for v in header) + " |",
        "| " + " | ".join("---" for _ in header) + " |",
    ]
    lines.extend("| " + " | ".join(esc(v) for v in row) + " |" for row in body)
    return Block(kind="table", text="\n".join(lines))


def iter_blocks(
    document: DocumentType, media_dir: Path, source_id: str
) -> Iterable[Block]:
    rels = document.part.rels
    for item in document.iter_inner_content():
        if isinstance(item, Paragraph):
            block = paragraph_block(item, rels, media_dir, source_id)
        elif isinstance(item, Table):
            block = table_block(item)
        else:
            block = None
        if block:
            yield block


def render_block(block: Block) -> str:
    chunks: list[str] = []
    if block.text:
        if block.level:
            chunks.append(f"{'#' * block.level} {block.text}")
        else:
            chunks.append(block.text)
    for image in block.images or []:
        chunks.append(f"![原文图片](assets/{image})")
    return "\n\n".join(chunks)


def import_doc(source_id: str, path: Path) -> None:
    if not path.exists():
        raise FileNotFoundError(path)
    out_dir = OUT_ROOT / source_id
    media_dir = out_dir / "assets"
    if out_dir.exists():
        shutil.rmtree(out_dir)
    media_dir.mkdir(parents=True)

    document = Document(path)
    blocks = list(iter_blocks(document, media_dir, source_id))
    title = clean(document.paragraphs[0].text) or path.stem

    sections: list[tuple[str, list[Block]]] = []
    current_name = "00-front-matter"
    current: list[Block] = []
    used: set[str] = set()

    for block in blocks:
        if block.level == 1 and block.text:
            if current:
                sections.append((current_name, current))
            base = f"{block.section or len(sections) + 1}-{slugify(block.text)}"
            name = base
            counter = 2
            while name in used:
                name = f"{base}-{counter}"
                counter += 1
            used.add(name)
            current_name = name
            current = [block]
        else:
            current.append(block)
    if current:
        sections.append((current_name, current))

    index_lines = [
        "---",
        f"title: {json.dumps(title, ensure_ascii=False)}",
        f"source_id: {source_id}",
        f"original_file: {json.dumps(path.name, ensure_ascii=False)}",
        "generated: true",
        "---",
        "",
        f"# {title}",
        "",
        "> 自动生成的来源层。请勿手工修改；规范知识应写入知识层目录。",
        "",
        "## 章节",
        "",
    ]
    manifest_sections = []
    for filename, section_blocks in sections:
        section_title = next(
            (b.text for b in section_blocks if b.level == 1 and b.text), filename
        )
        md_name = f"{filename}.md"
        content = [
            "---",
            f"source_id: {source_id}",
            f"source_file: {json.dumps(path.name, ensure_ascii=False)}",
            f"source_section: {json.dumps(section_title, ensure_ascii=False)}",
            "generated: true",
            "---",
            "",
        ]
        content.extend(render_block(block) for block in section_blocks)
        rendered = "\n\n".join(content).rstrip() + "\n"
        (out_dir / md_name).write_text(rendered, encoding="utf-8")
        index_lines.append(f"- [{section_title}]({md_name})")
        manifest_sections.append(
            {
                "file": md_name,
                "title": section_title,
                "blocks": len(section_blocks),
                "characters": sum(len(b.text) for b in section_blocks),
                "sha256": hashlib.sha256(rendered.encode("utf-8")).hexdigest(),
            }
        )

    (out_dir / "index.md").write_text("\n".join(index_lines) + "\n", encoding="utf-8")
    source_hash = hashlib.sha256(path.read_bytes()).hexdigest()
    manifest = {
        "source_id": source_id,
        "original_file": path.name,
        "sha256": source_hash,
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "images": len(list(media_dir.iterdir())),
        "sections": manifest_sections,
    }
    (out_dir / "manifest.json").write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
    )
    print(
        f"{source_id}: {len(sections)} sections, "
        f"{manifest['images']} images, {sum(x['characters'] for x in manifest_sections)} chars"
    )


def main() -> None:
    OUT_ROOT.mkdir(parents=True, exist_ok=True)
    for source_id, path in DOCS.items():
        import_doc(source_id, path)


if __name__ == "__main__":
    main()
